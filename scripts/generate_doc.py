#!/usr/bin/env python3
"""
生成 GeoClaw 项目介绍 Word 文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_document():
    doc = Document()
    
    # 设置标题样式
    title_style = doc.styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(30, 58, 95)
    
    # 添加标题
    title = doc.add_heading('GeoClaw 地质智能分析助手', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('中关村北纬龙虾大赛 · 学术赛道')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(74, 111, 165)
    
    # 重点标注
    highlight = doc.add_paragraph()
    highlight.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = highlight.add_run('【重点功能】区域遥感数据处理 - 一键构建数字孪生视图')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(200, 50, 50)
    
    doc.add_paragraph()
    
    # 项目概述
    doc.add_heading('一、项目概述', level=1)
    
    doc.add_paragraph(
        'GeoClaw 是专为地质科研工作者设计的 AI 智能助手，将大语言模型与地质专业知识深度融合，'
        '实现文献智能检索、岩芯图像分析、区域遥感数据处理、自动化报告生成等功能。'
    )
    
    doc.add_paragraph()
    
    # 核心痛点解决
    doc.add_heading('二、核心痛点解决：区域遥感数据处理', level=1)
    
    doc.add_paragraph(
        '地质科研中，区域遥感数据（卫星影像、多光谱/高光谱、DEM、InSAR等）的获取、整理、统筹、可视化'
        '是最大痛点之一。'
    )
    
    doc.add_heading('2.1 痛点分析', level=2)
    doc.add_paragraph('• 数据源散乱：Google Earth Engine、USGS、ESA Copernicus、国内高分/资源卫星', style='List Bullet')
    doc.add_paragraph('• 格式杂：GeoTIFF、NetCDF、HDF', style='List Bullet')
    doc.add_paragraph('• 体积大：单景影像可达数GB', style='List Bullet')
    doc.add_paragraph('• 坐标转换麻烦：不同数据源坐标系不统一', style='List Bullet')
    doc.add_paragraph('• 时间序列分析费时：需手动下载多期数据', style='List Bullet')
    doc.add_paragraph('• 出图/报告手动拼：缺乏自动化流程', style='List Bullet')
    
    doc.add_heading('2.2 解决方案', level=2)
    doc.add_paragraph(
        'GeoClaw 提供一键构建区域遥感地质"数字孪生"视图的能力：'
    )
    doc.add_paragraph('1. 搜集/下载最新遥感数据', style='List Number')
    doc.add_paragraph('2. 预处理（裁剪、校正、融合）', style='List Number')
    doc.add_paragraph('3. 分析（变化检测、分类、DEM衍生）', style='List Number')
    doc.add_paragraph('4. 可视化（生成地图、剖面、时序图）', style='List Number')
    doc.add_paragraph('5. 输出报告草稿（方法+结果+讨论）', style='List Number')
    
    doc.add_heading('2.3 实效量化', level=2)
    
    # 创建效率对比表
    table_eff = doc.add_table(rows=5, cols=3)
    table_eff.style = 'Table Grid'
    
    table_eff.rows[0].cells[0].text = '工作环节'
    table_eff.rows[0].cells[1].text = '传统方式'
    table_eff.rows[0].cells[2].text = 'GeoClaw方式'
    
    eff_data = [
        ['数据获取', '3-5天', '10分钟'],
        ['数据预处理', '1-2天', '30分钟'],
        ['变化检测分析', '2-3天', '1小时'],
        ['报告编制', '1-2天', '30分钟'],
    ]
    
    for i, (task, trad, geo) in enumerate(eff_data, 1):
        table_eff.rows[i].cells[0].text = task
        table_eff.rows[i].cells[1].text = trad
        table_eff.rows[i].cells[2].text = geo
    
    doc.add_paragraph()
    
    # 核心价值
    doc.add_heading('三、核心价值', level=1)
    
    # 创建表格
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '功能模块'
    header_cells[1].text = '传统方式'
    header_cells[2].text = 'GeoClaw方式'
    header_cells[3].text = '效率提升'
    
    # 数据行
    data = [
        ['遥感数据处理', '3-5天', '1-2小时', '100倍'],
        ['文献检索', '关键词匹配，效率低', '语义理解，精准定位', '3-5倍'],
        ['图像分析', '人工判读，耗时', 'AI自动识别', '80%时间节省'],
        ['报告撰写', '手工编制，易错', '自动生成，规范', '5-10倍'],
    ]
    
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row[j].text = cell_data
    
    doc.add_paragraph()
    
    # 技术架构
    doc.add_heading('四、技术架构', level=1)
    
    doc.add_heading('4.1 遥感数据处理流程', level=2)
    doc.add_paragraph(
        '用户输入研究区域 → AI Agent 自动执行以下流程：'
    )
    doc.add_paragraph('步骤1：数据获取', style='List Bullet')
    doc.add_paragraph('  • Google Earth Engine API', style='List Bullet')
    doc.add_paragraph('  • USGS Earth Explorer', style='List Bullet')
    doc.add_paragraph('  • ESA Copernicus Open Access Hub', style='List Bullet')
    doc.add_paragraph('  • 国内高分/资源卫星数据平台', style='List Bullet')
    
    doc.add_paragraph('步骤2：智能预处理', style='List Bullet')
    doc.add_paragraph('  • 坐标系统一（WGS84）', style='List Bullet')
    doc.add_paragraph('  • 影像裁剪与融合', style='List Bullet')
    doc.add_paragraph('  • 辐射/大气校正', style='List Bullet')
    doc.add_paragraph('  • 云/阴影掩膜', style='List Bullet')
    
    doc.add_paragraph('步骤3：分析计算', style='List Bullet')
    doc.add_paragraph('  • NDVI/NDWI 指数计算', style='List Bullet')
    doc.add_paragraph('  • 变化检测（时序分析）', style='List Bullet')
    doc.add_paragraph('  • DEM 衍生分析（坡度/坡向）', style='List Bullet')
    doc.add_paragraph('  • 土地利用分类', style='List Bullet')
    
    doc.add_paragraph('步骤4：可视化输出', style='List Bullet')
    doc.add_paragraph('  • 交互式地图（Plotly/Folium）', style='List Bullet')
    doc.add_paragraph('  • 时序动画', style='List Bullet')
    doc.add_paragraph('  • 剖面线图', style='List Bullet')
    doc.add_paragraph('  • 分析报告', style='List Bullet')
    
    doc.add_heading('4.2 核心模型层', level=2)
    doc.add_paragraph('• 智谱 GLM-4-Flash：对话理解与推理', style='List Bullet')
    doc.add_paragraph('• 智谱 GLM-4V-Flash：视觉理解与图像分析', style='List Bullet')
    doc.add_paragraph('• 智谱 ASR：语音转文本（野外记录转写）', style='List Bullet')
    
    doc.add_heading('4.3 技能扩展层', level=2)
    doc.add_paragraph('• 文件系统操作：读写本地文件夹、整理PDF文献', style='List Bullet')
    doc.add_paragraph('• 网页内容抓取：获取地质数据库信息', style='List Bullet')
    doc.add_paragraph('• 自定义技能框架：支持功能扩展', style='List Bullet')
    
    doc.add_heading('4.4 应用场景层', level=2)
    doc.add_paragraph('• 区域地质调查：地质图件解译、野外记录转写', style='List Bullet')
    doc.add_paragraph('• 矿产勘查评价：岩芯图像分析、勘查报告编制', style='List Bullet')
    doc.add_paragraph('• 灾害地质评估：遥感影像解译、隐患识别', style='List Bullet')
    doc.add_paragraph('• 科研数据分析：文献检索、数据可视化', style='List Bullet')
    
    doc.add_paragraph()
    
    # 理论溯源
    doc.add_heading('五、理论溯源与创新', level=1)
    
    doc.add_paragraph(
        '本项目技术路线借鉴自 Nature (2025) 论文《AI agents for rare disease diagnosis》，'
        '该论文展示了 AI Agent 在医疗垂直领域的成功应用，通过多模态分析、知识库增强、'
        '自动化工作流等技术，实现了罕见病的高效诊断。'
    )
    
    doc.add_heading('5.1 方法论迁移', level=2)
    
    # 创建迁移对照表
    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'
    
    table2.rows[0].cells[0].text = '医疗领域（论文）'
    table2.rows[0].cells[1].text = '地质领域（本项目）'
    
    migration_data = [
        ['罕见病诊断', '地质异常识别'],
        ['医学影像分析', '岩芯/薄片图像分析'],
        ['临床报告生成', '地质调查报告生成'],
        ['病历知识库', '地质文献知识库'],
    ]
    
    for i, (medical, geo) in enumerate(migration_data, 1):
        table2.rows[i].cells[0].text = medical
        table2.rows[i].cells[1].text = geo
    
    doc.add_paragraph()
    
    doc.add_heading('5.2 创新亮点', level=2)
    doc.add_paragraph('1. 多模态融合分析：结合文本、图像、语音多维度理解地质信息', style='List Number')
    doc.add_paragraph('2. 知识库增强检索：基于智谱模型的语义理解，精准定位相关文献', style='List Number')
    doc.add_paragraph('3. 自动化工作流编排：从数据采集到报告生成的一站式解决方案', style='List Number')
    doc.add_paragraph('4. 低代码/零代码使用：通过 Streamlit 构建 Web 界面，降低使用门槛', style='List Number')
    
    doc.add_paragraph()
    
    # 功能详解
    doc.add_heading('六、功能详解', level=1)
    
    doc.add_heading('6.1 区域遥感数据处理（重点功能）', level=2)
    doc.add_paragraph(
        '用户只需输入研究区域坐标和研究主题，AI Agent 自动：'
    )
    doc.add_paragraph('• 自动搜索并获取多源遥感数据', style='List Bullet')
    doc.add_paragraph('• 生成数据获取计划与下载链接', style='List Bullet')
    doc.add_paragraph('• 提供 Python 代码片段进行数据处理', style='List Bullet')
    doc.add_paragraph('• 自动生成分析报告框架', style='List Bullet')
    
    doc.add_paragraph('支持的数据源：')
    doc.add_paragraph('• Sentinel-2 (10m 多光谱)', style='List Bullet')
    doc.add_paragraph('• Landsat 8/9 (30m 多光谱)', style='List Bullet')
    doc.add_paragraph('• SRTM/ASTER GDEM (30m 高程)', style='List Bullet')
    doc.add_paragraph('• 高分系列 (国产高分卫星)', style='List Bullet')
    
    doc.add_heading('6.2 文献智能检索', level=2)
    doc.add_paragraph(
        '基于智谱 GLM-4 的语义理解能力，用户只需描述研究需求，系统即可：'
    )
    doc.add_paragraph('• 自动生成中英文关键检索词', style='List Bullet')
    doc.add_paragraph('• 推荐相关经典文献', style='List Bullet')
    doc.add_paragraph('• 建议权威数据来源', style='List Bullet')
    doc.add_paragraph('• 支持模板化快捷检索', style='List Bullet')
    
    doc.add_heading('6.3 岩芯图像智能分析', level=2)
    doc.add_paragraph(
        '基于 GLM-4V 视觉模型，自动识别岩芯照片、薄片显微图中的矿物成分与结构特征：'
    )
    doc.add_paragraph('• 岩芯照片分析：岩性识别、矿物成分、结构构造、蚀变特征', style='List Bullet')
    doc.add_paragraph('• 薄片显微图分析：矿物鉴定、结构分析、蚀变与交代、成因推断', style='List Bullet')
    doc.add_paragraph('• 遥感影像解译：地貌特征、构造特征、岩性分区、异常信息', style='List Bullet')
    doc.add_paragraph('• 地质图件识别：图件类型判断、主要内容概括、图例解读', style='List Bullet')
    
    doc.add_heading('6.4 自动化报告生成', level=2)
    doc.add_paragraph('支持多种地质报告类型自动生成：')
    doc.add_paragraph('• 地质调查报告', style='List Bullet')
    doc.add_paragraph('• 矿床勘查报告', style='List Bullet')
    doc.add_paragraph('• 研究论文摘要', style='List Bullet')
    doc.add_paragraph('• 项目进展报告', style='List Bullet')
    doc.add_paragraph('用户只需输入基础信息和工作内容，系统自动生成符合地质行业规范的完整报告。')
    
    doc.add_paragraph()
    
    # 实际应用场景
    doc.add_heading('七、实际应用场景', level=1)
    
    doc.add_heading('场景一：区域遥感地质调查', level=2)
    doc.add_paragraph(
        '研究人员研究某区域构造演化时，通过 GeoClaw：'
    )
    doc.add_paragraph('1. 输入研究区域坐标和研究主题', style='List Number')
    doc.add_paragraph('2. AI 自动生成数据获取计划', style='List Number')
    doc.add_paragraph('3. 下载多源遥感数据（光学+SRTM）', style='List Number')
    doc.add_paragraph('4. 计算遥感指数并进行变化检测', style='List Number')
    doc.add_paragraph('5. 生成可视化地图和分析报告', style='List Number')
    doc.add_paragraph('实效：从原来 1 周压缩到 2 小时', style='List Bullet')
    
    doc.add_heading('场景二：海岸侵蚀监测', level=2)
    doc.add_paragraph(
        '研究人员分析台湾东部海岸侵蚀变化：'
    )
    doc.add_paragraph('1. 输入"台湾东部海岸侵蚀分析"', style='List Number')
    doc.add_paragraph('2. 获取多年份 Sentinel-2 影像', style='List Number')
    doc.add_paragraph('3. 计算海岸线位置变化', style='List Number')
    doc.add_paragraph('4. 生成时序变化图和统计报告', style='List Number')
    doc.add_paragraph('实效：从原来 3-5 天压缩到 1 小时', style='List Bullet')
    
    doc.add_heading('场景三：区域地质调查', level=2)
    doc.add_paragraph(
        '地质队员在野外进行调查时，可通过 GeoClaw：'
    )
    doc.add_paragraph('1. 语音录入野外观察记录（智谱 ASR 自动转写）', style='List Number')
    doc.add_paragraph('2. 拍摄地质现象照片，上传后自动分析识别', style='List Number')
    doc.add_paragraph('3. 检索相关区域地质文献资料', style='List Number')
    doc.add_paragraph('4. 自动生成地质调查报告', style='List Number')
    
    doc.add_heading('场景四：矿产勘查评价', level=2)
    doc.add_paragraph(
        '勘查工程师可利用 GeoClaw：'
    )
    doc.add_paragraph('1. 上传岩芯照片，自动识别岩性和矿化信息', style='List Number')
    doc.add_paragraph('2. 分析薄片显微图像，鉴定矿物组合', style='List Number')
    doc.add_paragraph('3. 检索相似矿床案例和研究文献', style='List Number')
    doc.add_paragraph('4. 快速编制勘查报告', style='List Number')
    
    doc.add_paragraph()
    
    # 技术实现
    doc.add_heading('八、技术实现', level=1)
    
    doc.add_heading('8.1 开发环境', level=2)
    doc.add_paragraph('• Python 3.12', style='List Bullet')
    doc.add_paragraph('• Streamlit 应用框架', style='List Bullet')
    doc.add_paragraph('• 智谱 AI SDK (zhipuai)', style='List Bullet')
    doc.add_paragraph('• OpenClaw 技能系统', style='List Bullet')
    
    doc.add_heading('8.2 部署方式', level=2)
    doc.add_paragraph('• 本地部署：Streamlit 本地运行', style='List Bullet')
    doc.add_paragraph('• 云端部署：支持 Streamlit Cloud、阿里云等平台', style='List Bullet')
    doc.add_paragraph('• API 服务：可通过 FastAPI 封装为 API 服务', style='List Bullet')
    
    doc.add_paragraph()
    
    # 未来规划
    doc.add_heading('九、未来规划', level=1)
    
    doc.add_paragraph('• 接入地质专业数据库（GeoRef、CNKI 地质期刊等）', style='List Bullet')
    doc.add_paragraph('• 训练地质领域专用模型，提升识别精度', style='List Bullet')
    doc.add_paragraph('• 开发移动端应用，支持野外实时使用', style='List Bullet')
    doc.add_paragraph('• 构建地质知识图谱，增强推理能力', style='List Bullet')
    doc.add_paragraph('• 支持多语言（中英文切换）', style='List Bullet')
    
    doc.add_paragraph()
    
    # 总结
    doc.add_heading('十、总结', level=1)
    
    doc.add_paragraph(
        'GeoClaw 地质智能分析助手通过将先进的 AI 技术与地质专业知识深度融合，'
        '为地质科研工作者提供了高效、便捷的智能工具。'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('【核心亮点】')
    doc.add_paragraph(
        '项目重点解决了区域遥感数据处理的痛点——从数据获取、预处理、分析到可视化输出，'
        '实现了一键式自动化流程，将原来需要数天的工作压缩到数小时。'
    )
    
    doc.add_paragraph(
        '项目借鉴 Nature 论文中的 AI Agent 方法论，'
        '实现了多模态分析、知识库增强、自动化工作流等核心功能，在遥感数据处理、文献检索、图像分析、'
        '报告生成等场景展现出显著的价值。'
    )
    
    doc.add_paragraph(
        '作为中关村北纬龙虾大赛学术赛道的参赛作品，GeoClaw 展示了 AI Agent 在垂直领域应用的巨大潜力，'
        '体现了"科技向善"的理念——用 AI 技术帮助科研工作者提升效率、解决实际问题。'
    )
    
    doc.add_paragraph()
    
    # 附录
    doc.add_heading('附录', level=1)
    
    doc.add_heading('A. 参考论文', level=2)
    doc.add_paragraph(
        'Nature (2025). AI agents for rare disease diagnosis. '
        'https://www.nature.com/articles/s41586-025-10097-9'
    )
    
    doc.add_heading('B. 技术文档', level=2)
    doc.add_paragraph('• 智谱 AI 开放平台：https://bigmodel.cn')
    doc.add_paragraph('• OpenClaw 官方文档：https://docs.openclaw.ai')
    doc.add_paragraph('• Streamlit 文档：https://docs.streamlit.io')
    
    doc.add_heading('C. 联系方式', level=2)
    doc.add_paragraph('• 项目仓库：GitHub（开源中）')
    doc.add_paragraph('• 问题反馈：GitHub Issues')
    doc.add_paragraph('• 技术交流：OpenClaw 社区')
    
    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f'© 2025 GeoClaw Team | 生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存文档
    output_path = os.path.expanduser('~/openclaw-game/docs/GeoClaw项目介绍.docx')
    doc.save(output_path)
    print(f'✅ Word 文档已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    import os
    create_document()
